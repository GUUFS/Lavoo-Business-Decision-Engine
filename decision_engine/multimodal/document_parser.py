"""
Document Parser Module

Handles extraction of text from various document formats (PDF, DOCX, TXT).
Uses free, open-source libraries to parse business documents.

Example Use Cases:
    - Extract text from business reports (PDF)
    - Parse financial statements (DOCX)
    - Read plain text files with business data
    - Process uploaded documentation
"""

import logging
from pathlib import Path
from typing import Dict, Optional, Union
from io import BytesIO

# PDF parsing libraries
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logging.warning("PyPDF2 not installed. Install with: uv add PyPDF2")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    logging.warning("pdfplumber not installed. Install with: uv add pdfplumber")

# DOCX parsing
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not installed. Install with: uv add python-docx")

from .config import (
    SUPPORTED_DOC_FORMATS,
    MAX_FILE_SIZE_MB,
    EXTRACTION_CONFIG,
)

logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parses documents (PDF, DOCX, TXT) to extract text content.
    
    Supports:
        - PDF: Using PyPDF2 or pdfplumber
        - DOCX: Using python-docx
        - TXT: Native Python
    """
    
    def __init__(self, prefer_pdfplumber: bool = True):
        """
        Initialize DocumentParser.
        
        Args:
            prefer_pdfplumber (bool): Prefer pdfplumber over PyPDF2 for PDFs
                                     (pdfplumber is more accurate but slower)
        """
        self.prefer_pdfplumber = prefer_pdfplumber and PDFPLUMBER_AVAILABLE
        
        # Check library availability
        if not PYPDF2_AVAILABLE and not PDFPLUMBER_AVAILABLE:
            logger.warning("No PDF parsing library available. Install PyPDF2 or pdfplumber")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX parsing disabled")
        
        logger.info(f"DocumentParser initialized (PDF parser: {'pdfplumber' if self.prefer_pdfplumber else 'PyPDF2'})")
    
    def _validate_document_file(self, file_path: Path) -> bool:
        """
        Validate document file format and size.
        
        Args:
            file_path (Path): Path to document file
            
        Returns:
            bool: True if valid, False otherwise
        """
        # Check file extension
        if file_path.suffix.lower() not in SUPPORTED_DOC_FORMATS:
            logger.error(f"Unsupported document format: {file_path.suffix}")
            return False
        
        # Check file size
        file_size_mb = file_path.stat().st_size / (1024 * 1024)
        if file_size_mb > MAX_FILE_SIZE_MB:
            logger.error(f"File too large: {file_size_mb:.2f}MB (max: {MAX_FILE_SIZE_MB}MB)")
            return False
        
        logger.debug(f"Document validated: {file_path.name} ({file_size_mb:.2f}MB)")
        return True
    
    def extract_text_from_file(self, file_path: Union[str, Path]) -> Dict[str, Union[str, bool]]:
        """
        Extract text from a document file.
        
        Args:
            file_path (str | Path): Path to document file
            
        Returns:
            dict: {
                'success': bool,
                'text': str (extracted text),
                'error': str (error message if failed),
                'metadata': dict (page count, format, etc.)
            }
        """
        try:
            file_path = Path(file_path)
            
            # Validate file
            if not file_path.exists():
                return {"success": False, "text": "", "error": "File not found"}
            
            if not self._validate_document_file(file_path):
                return {"success": False, "text": "", "error": "Invalid document file"}
            
            # Route to appropriate parser based on file extension
            suffix = file_path.suffix.lower()
            logger.info(f"Processing {suffix} file: {file_path.name}")
            
            if suffix == ".pdf":
                return self._extract_from_pdf(file_path)
            elif suffix == ".docx":
                return self._extract_from_docx(file_path)
            elif suffix == ".txt":
                return self._extract_from_txt(file_path)
            else:
                return {"success": False, "text": "", "error": f"Unsupported format: {suffix}"}
        
        except Exception as e:
            logger.error(f"Error extracting text from document: {e}", exc_info=True)
            return {
                "success": False,
                "text": "",
                "error": f"Document parsing failed: {str(e)}",
                "metadata": {}
            }
    
    def extract_text_from_bytes(
        self, 
        doc_bytes: bytes, 
        filename: str = "unknown.pdf"
    ) -> Dict[str, Union[str, bool]]:
        """
        Extract text from document bytes (useful for uploaded files).
        
        Args:
            doc_bytes (bytes): Document data as bytes
            filename (str): Original filename (determines parsing method)
            
        Returns:
            dict: Same format as extract_text_from_file()
        """
        try:
            # Check size
            size_mb = len(doc_bytes) / (1024 * 1024)
            if size_mb > MAX_FILE_SIZE_MB:
                return {
                    "success": False,
                    "text": "",
                    "error": f"File too large: {size_mb:.2f}MB (max: {MAX_FILE_SIZE_MB}MB)"
                }
            
            # Determine format from filename
            suffix = Path(filename).suffix.lower()
            logger.info(f"Processing {suffix} bytes: {filename}")
            
            if suffix == ".pdf":
                return self._extract_from_pdf_bytes(doc_bytes, filename)
            elif suffix == ".docx":
                return self._extract_from_docx_bytes(doc_bytes, filename)
            elif suffix == ".txt":
                return self._extract_from_txt_bytes(doc_bytes, filename)
            else:
                return {"success": False, "text": "", "error": f"Unsupported format: {suffix}"}
        
        except Exception as e:
            logger.error(f"Error extracting text from bytes: {e}", exc_info=True)
            return {
                "success": False,
                "text": "",
                "error": f"Document parsing failed: {str(e)}",
                "metadata": {}
            }
    
    def _extract_from_pdf(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Extract text from PDF file."""
        # Try pdfplumber first (more accurate)
        if self.prefer_pdfplumber:
            try:
                return self._extract_with_pdfplumber(file_path)
            except Exception as e:
                logger.warning(f"pdfplumber failed, falling back to PyPDF2: {e}")
        
        # Fallback to PyPDF2
        if PYPDF2_AVAILABLE:
            return self._extract_with_pypdf2(file_path)
        
        return {"success": False, "text": "", "error": "No PDF parser available"}
    
    def _extract_with_pypdf2(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Extract text using PyPDF2."""
        text_parts = []
        page_count = 0
        
        with open(file_path, "rb") as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            page_count = len(reader.pages)
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                logger.debug(f"Extracted page {page_num}/{page_count}")
        
        full_text = "\n\n".join(text_parts).strip()
        
        return self._finalize_extraction(
            full_text, 
            {"page_count": page_count, "parser": "PyPDF2", "filename": file_path.name}
        )
    
    def _extract_with_pdfplumber(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Extract text using pdfplumber (more accurate)."""
        text_parts = []
        page_count = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                logger.debug(f"Extracted page {page_num}/{page_count}")
        
        full_text = "\n\n".join(text_parts).strip()
        
        return self._finalize_extraction(
            full_text,
            {"page_count": page_count, "parser": "pdfplumber", "filename": file_path.name}
        )
    
    def _extract_from_pdf_bytes(self, pdf_bytes: bytes, filename: str) -> Dict[str, Union[str, bool]]:
        """Extract text from PDF bytes."""
        if self.prefer_pdfplumber:
            try:
                with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
                    text_parts = [page.extract_text() for page in pdf.pages if page.extract_text()]
                    full_text = "\n\n".join(text_parts).strip()
                    return self._finalize_extraction(
                        full_text,
                        {"page_count": len(pdf.pages), "parser": "pdfplumber", "filename": filename}
                    )
            except Exception as e:
                logger.warning(f"pdfplumber failed on bytes, trying PyPDF2: {e}")
        
        if PYPDF2_AVAILABLE:
            reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
            text_parts = [page.extract_text() for page in reader.pages if page.extract_text()]
            full_text = "\n\n".join(text_parts).strip()
            return self._finalize_extraction(
                full_text,
                {"page_count": len(reader.pages), "parser": "PyPDF2", "filename": filename}
            )
        
        return {"success": False, "text": "", "error": "No PDF parser available"}
    
    def _extract_from_docx(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            return {"success": False, "text": "", "error": "python-docx not installed"}
        
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs).strip()
        
        return self._finalize_extraction(
            full_text,
            {"paragraph_count": len(paragraphs), "parser": "python-docx", "filename": file_path.name}
        )
    
    def _extract_from_docx_bytes(self, docx_bytes: bytes, filename: str) -> Dict[str, Union[str, bool]]:
        """Extract text from DOCX bytes."""
        if not DOCX_AVAILABLE:
            return {"success": False, "text": "", "error": "python-docx not installed"}
        
        doc = docx.Document(BytesIO(docx_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        full_text = "\n\n".join(paragraphs).strip()
        
        return self._finalize_extraction(
            full_text,
            {"paragraph_count": len(paragraphs), "parser": "python-docx", "filename": filename}
        )
    
    def _extract_from_txt(self, file_path: Path) -> Dict[str, Union[str, bool]]:
        """Extract text from TXT file."""
        try:
            # Try UTF-8 first
            with open(file_path, "r", encoding="utf-8") as f:
                text = f.read().strip()
        except UnicodeDecodeError:
            # Fallback to latin-1
            logger.warning(f"UTF-8 decode failed for {file_path.name}, trying latin-1")
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read().strip()
        
        return self._finalize_extraction(
            text,
            {"encoding": "utf-8/latin-1", "parser": "native", "filename": file_path.name}
        )
    
    def _extract_from_txt_bytes(self, txt_bytes: bytes, filename: str) -> Dict[str, Union[str, bool]]:
        """Extract text from TXT bytes."""
        try:
            text = txt_bytes.decode("utf-8").strip()
        except UnicodeDecodeError:
            logger.warning(f"UTF-8 decode failed for {filename}, trying latin-1")
            text = txt_bytes.decode("latin-1").strip()
        
        return self._finalize_extraction(
            text,
            {"encoding": "utf-8/latin-1", "parser": "native", "filename": filename}
        )
    
    def _finalize_extraction(self, text: str, metadata: dict) -> Dict[str, Union[str, bool]]:
        """
        Finalize extraction: validate length and truncate if needed.
        
        Args:
            text (str): Extracted text
            metadata (dict): Document metadata
            
        Returns:
            dict: Standardized extraction result
        """
        # Check minimum length
        if len(text) < EXTRACTION_CONFIG["min_text_length"]:
            logger.warning(f"Extracted text too short: {len(text)} chars")
            return {
                "success": True,
                "text": text,
                "error": "No significant text found in document",
                "metadata": {**metadata, "char_count": len(text)}
            }
        
        # Truncate if too long
        if len(text) > EXTRACTION_CONFIG["max_text_length"]:
            logger.warning(f"Truncating text from {len(text)} to {EXTRACTION_CONFIG['max_text_length']} chars")
            text = text[:EXTRACTION_CONFIG["max_text_length"]] + "\n\n... [truncated]"
        
        logger.info(f"Successfully extracted {len(text)} characters from document")
        
        return {
            "success": True,
            "text": text,
            "error": None,
            "metadata": {**metadata, "char_count": len(text)}
        }
